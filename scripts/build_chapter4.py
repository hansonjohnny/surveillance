"""
Builds Chapter 4 (System Development and Testing) and appends it to a copy of
Surveillance_AI_Chapters_1_to_3.docx, inserting the new chapter directly
before the existing "REFERENCES" heading so the final order is:
Chapter 1 -> Chapter 2 -> Chapter 3 -> Chapter 4 -> References.

Run with: cmd /c "python scripts\\build_chapter4.py"
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = "Surveillance_AI_Chapters_1_to_3.docx"
OUT = "Surveillance_AI_Chapters_1_to_4.docx"

doc = Document(SRC)

# ---------------------------------------------------------------------------
# Locate the "REFERENCES" heading so new content can be inserted before it.
# ---------------------------------------------------------------------------
references_para = None
for p in doc.paragraphs:
    if p.text.strip() == "REFERENCES" and p.style.name == "Heading 1":
        references_para = p
        break

if references_para is None:
    raise RuntimeError('Could not find a "REFERENCES" Heading 1 paragraph.')

# ---------------------------------------------------------------------------
# Formatting helpers replicating the exact styles/paragraph formats already
# used throughout the document (inspected from Chapters One and Two).
# ---------------------------------------------------------------------------

def _new(style):
    p = references_para.insert_paragraph_before("", style=style)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return p


def h1(text):
    p = _new("Heading 1")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(text)
    return p


def h2(text):
    p = _new("Heading 2")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(text)
    return p


def h3(text):
    p = _new("Heading 3")
    p.add_run(text)
    return p


def body(text):
    p = _new("Normal")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(10)
    p.add_run(text)
    return p


def bullet_list(items):
    for item in items:
        p = _new("List Paragraph")
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        if isinstance(item, tuple):
            lead, rest = item
            r1 = p.add_run(f"\u2022\t{lead} ")
            r1.bold = True
            p.add_run(rest)
        else:
            p.add_run(f"\u2022\t{item}")


def numbered_list(items):
    for i, item in enumerate(items, start=1):
        p = _new("List Paragraph")
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        if isinstance(item, tuple):
            lead, rest = item
            r1 = p.add_run(f"{i}.\t{lead} ")
            r1.bold = True
            p.add_run(rest)
        else:
            p.add_run(f"{i}.\t{item}")


def code_block(lines):
    """Monospace, left-aligned, single-spaced code listing."""
    for line in lines:
        p = _new("Normal")
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.3)
        r = p.add_run(line if line else "\u00A0")
        r.font.name = "Consolas"
        r.font.size = Pt(9)


def figure_caption(number, caption, space_after=14, label="Figure"):
    cap = _new("Normal")
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(space_after)
    r = cap.add_run(f"{label} {number}: {caption}")
    r.italic = True


def _set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "999999")
        borders.append(el)
    tblPr.append(borders)


def insert_table_before(rows, cols, widths):
    table = doc.add_table(rows=rows, cols=cols)
    _set_table_borders(table)
    for i, w in enumerate(widths):
        for row in table.rows:
            row.cells[i].width = Inches(w)
    # Relocate the freshly-appended table so it sits directly before
    # the REFERENCES heading instead of at the very end of the document.
    references_para._p.addprevious(table._tbl)
    return table


def fill_cell(cell, text, bold=False, size=9, center=False):
    cell.paragraphs[0].text = ""
    r = cell.paragraphs[0].add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    if center:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.paragraphs[0].paragraph_format.space_after = Pt(4)


# ---------------------------------------------------------------------------
# CHAPTER FOUR
# ---------------------------------------------------------------------------

h1("CHAPTER FOUR")
h2("SYSTEM DEVELOPMENT AND TESTING")

body(
    "This chapter describes how the design presented in Chapter Three was "
    "implemented and verified. It covers the development approach and "
    "environment used to build the system, the organisation and key "
    "mechanisms of the codebase, the application's main user interfaces, the "
    "testing carried out on the completed system, and the documentation "
    "produced alongside it."
)

h3("4.1 Implementing of the Design")
body(
    "The system was implemented feature by feature rather than as a single "
    "monolithic build: authentication and the Supabase schema (with Row "
    "Level Security enabled on every table) were established first, "
    "followed by the core monitoring loop and its sensor integrations, then "
    "the three-channel alert pipeline, the five main application screens, "
    "the tiered subscription and administrator model, and finally the "
    "manual test pass described in Section 4.2. For each feature, the "
    "smallest working version was built first and refactored only once "
    "genuine repetition or complexity appeared, in keeping with the "
    "project's own development conventions, which explicitly favour "
    "readable, minimal implementations over premature abstraction."
)
body(
    "Development was carried out in Visual Studio Code against an Expo "
    "SDK 56 / React Native project written entirely in TypeScript. Because "
    "the system depends on native modules (camera, microphone, background "
    "tasks, accelerometer) that are unavailable in a plain Expo Go sandbox, "
    "a custom Expo Dev Client was built and distributed through Expo "
    "Application Services (EAS) so that the application could be installed "
    "and iterated on directly on a physical Android device, which the "
    "project's own testing documentation identifies as mandatory since the "
    "simulator cannot exercise camera, microphone, accelerometer, "
    "background execution, or push notifications. On the backend, the "
    "Supabase CLI was used to author and apply PostgreSQL migrations and to "
    "deploy the seven Edge Functions, and static analysis was enforced "
    "through the project's `expo lint` script."
)

h3("4.2 Programming and Coding")
body(
    "The codebase is organised so that each concern has a single, "
    "predictable home: screens live under /app (grouped into authentication, "
    "onboarding, and tab routes managed by Expo Router), presentational "
    "components live under /components, framework-independent business "
    "logic lives under /lib (monitoring.ts, vision.ts, audio.ts, alerts.ts, "
    "location.ts, sensors.ts, plans.ts, prompts.ts), Zustand state "
    "containers live under /store, and the seven serverless functions that "
    "hold third-party credentials live under /supabase/functions. This "
    "separation keeps the AI prompts, in prompts.ts, and the plan "
    "parameters, in plans.ts, each centralised in one file, so that "
    "adjusting a prompt or a subscription tier does not require touching "
    "the screens or state containers that use them."
)
body(
    "The monitoring orchestrator in monitoring.ts is the most important "
    "module in the client, since it implements the passive capture cycle "
    "described in Chapter Three. Figure 4.1 reproduces its core logic: the "
    "camera snapshot and GPS fix are requested together with Promise.all so "
    "that neither sensor waits on the other, the resulting event is written "
    "to the Event Log immediately with a placeholder summary so the "
    "interface never appears to hang, and the image-based and audio-based "
    "risk levels returned by the two Edge Functions are later reconciled by "
    "combineRisks, which simply keeps whichever of the two carries the "
    "higher severity."
)
code_block([
    "function combineRisks(a: RiskLevel, b: RiskLevel | null): RiskLevel {",
    "  if (!b) return a;",
    "  const order: Record<RiskLevel, number> = { low: 0, medium: 1, high: 2 };",
    "  return order[a] >= order[b] ? a : b;",
    "}",
    "",
    "// Snapshot and GPS captured in parallel \u2014 neither blocks the other.",
    "const [photoUri, location] = await Promise.all([",
    "  takeSnapshot(),",
    "  getCurrentLocation(),",
    "]);",
    "",
    "// Event is logged immediately with a placeholder summary; the AI",
    "// result patches it in place once analyse-image resolves.",
    "useAlertStore.getState().addEvent(event);",
])
figure_caption("4.1", "Parallel sensor capture and risk-combination logic from monitoring.ts.")

body(
    "On the server side, the daily artificial-intelligence usage cap "
    "described in Chapter Three is enforced by a single atomic PostgreSQL "
    "function, increment_ai_usage, rather than by a separate read followed "
    "by a write from the Edge Function. Figure 4.2 shows the relevant "
    "clause: the INSERT ... ON CONFLICT DO UPDATE only increments "
    "call_count while it remains below the caller's cap, so two concurrent "
    "monitoring cycles for the same user on the same day cannot both "
    "increment past the limit."
)
code_block([
    "INSERT INTO ai_usage (user_id, date, call_count)",
    "VALUES (p_user_id, p_date, 1)",
    "ON CONFLICT (user_id, date) DO UPDATE",
    "  SET call_count = ai_usage.call_count + 1",
    "  WHERE ai_usage.call_count < p_cap",
    "RETURNING ai_usage.call_count INTO v_count;",
])
figure_caption("4.2", "Atomic, race-condition-safe usage-cap increment (increment_ai_usage RPC).")

h3("4.3 Main User Interfaces")
body(
    "The application exposes five main screens behind a bottom tab "
    "navigator, reached only after a user has signed in and completed the "
    "twelve-screen onboarding sequence described in Chapter Three."
)
bullet_list([
    ("Home (home.tsx).", "shows the single Start / Stop Surveillance control, an animated shield status indicator, the current risk badge, the most recent AI summary, a manual SOS button, and, where a permission has been denied, a degraded-mode banner."),
    ("Live (live.tsx).", "renders LiveMap.tsx, a WebView hosting a self-contained Leaflet page with dark CartoDB tiles, showing the user's current position, alongside the active risk badge and latest cycle summary."),
    ("Log (log.tsx).", "lists every event recorded during past and current sessions using EventCard.tsx, each showing a timestamp, risk badge, and AI summary; tapping a card opens ExpandedEventCard.tsx with the full report, photo, and transcript."),
    ("Alerts (alerts.tsx).", "lists every High-risk event that triggered the alert pipeline using AlertCard.tsx, showing which channels (SMS, email, call) were sent, to whom, and what the AI detected."),
    ("Settings (settings.tsx).", "exposes the emergency contact fields, MonitoringIntervalPicker.tsx and ShakeSensitivityPicker.tsx, the stealth-mode and wellness check-in configuration, the user's current subscription plan and usage, and a sign-out control."),
])
body(
    "When stealth mode is enabled, StealthOverlay.tsx dims the screen to "
    "black immediately after a session starts, and a single tap wakes it "
    "for approximately three seconds before it fades back to black, so that "
    "the session continues to run without any UI visible to a bystander. "
    "Device screenshots of these screens are provided in Appendix B."
)

h3("4.4 Testing of the New System")
body(
    "Because the system's most safety-critical behaviour, camera capture, "
    "microphone recording, accelerometer-based shake detection, background "
    "execution, and push notifications, cannot be exercised in a simulator, "
    "the system was tested manually on a physical Android device running "
    "the custom Expo Dev Client described in Section 4.1. Testing followed "
    "a black-box, scenario-based checklist covering every functional "
    "requirement in Section 3.2.1, summarised in Table 4.1."
)

widths = [1.3, 2.3, 2.4, 0.9]
table = insert_table_before(19, 4, widths)
headers = ["Test Area", "Test Case", "Expected Result", "Status"]
for i, htext in enumerate(headers):
    fill_cell(table.rows[0].cells[i], htext, bold=True, center=(i == 3))
    _set_cell_shading(table.rows[0].cells[i], "D9D9D9")

rows_data = [
    ("Authentication", "Sign up, sign in, wrong-password handling, password reset via deep link", "Account created/authenticated; invalid credentials rejected without crashing; reset link reopens app and updates password", "Pass"),
    ("Onboarding", "Complete all 12 onboarding screens and enter emergency contact", "Contact details persist and pre-fill Settings; onboarding does not replay after force-quit", "Pass"),
    ("Session Control", "Tap Start Surveillance, then Stop Surveillance", "Session record created/closed; timer runs; screen stays awake while active", "Pass"),
    ("Sensor Capture & AI Analysis", "Run a session and observe the Event Log", "Event appears immediately with placeholder summary, then updates with risk level, AI summary, and real GPS coordinates", "Pass"),
    ("Background Execution", "Minimise the app and lock the screen during a session", "Monitoring cycles continue; new events appear after reopening the app", "Pass"),
    ("Shake Detection", "Shake the device firmly while a session is active", "Immediate High-risk event logged with source \"shake\"; repeat shake within 5 seconds does not duplicate the alert", "Pass"),
    ("Medium-Risk Notification", "Mock a medium-risk AI response", "Local push notification delivered; no SMS or email sent", "Pass"),
    ("High-Risk Alert \u2013 Email", "Mock a high-risk AI response with a contact configured", "Full report email received via SendGrid with AI summary and GPS link", "Pass"),
    ("High-Risk Alert \u2013 SMS", "Mock a high-risk AI response with a contact configured", "SMS received via Twilio with GPS link and AI summary", "Pending \u2013 Twilio credentials not yet provisioned"),
    ("High-Risk Alert \u2013 Voice Call", "Shake event combined with a High AI score", "Automated Twilio voice call places to the emergency contact", "Pending \u2013 Twilio credentials not yet provisioned"),
    ("Alert Deduplication", "Trigger High risk twice for the same event", "Only one alert record is created; channels are not fired twice", "Pass"),
    ("Stealth Mode", "Enable stealth mode and start a session", "Screen fades to black; a single tap reveals the UI for ~3 seconds before dimming again", "Pass"),
    ("Wellness Check-In", "Configure a check-in time and let it lapse without confirming", "Emergency contact alerted after a 10-minute grace period; confirming \"I'm Safe\" cancels the alert", "Pass"),
    ("Live Map", "Move the device during an active session", "GPS marker on the Live screen updates to follow the device's real position", "Pass"),
    ("Settings Persistence", "Change contact details, interval, and sensitivity, then force-quit", "All changed values remain after reopening; the new interval is used on the next session", "Pass"),
    ("Degraded Mode", "Deny camera and/or microphone permission", "Amber banner shown; monitoring continues with the remaining sensors; app does not crash", "Pass"),
    ("Offline / Edge Cases", "Start a session with no internet connection or revoked location permission", "Cycles continue; events log a safe fallback summary; no alert fires; app does not crash", "Pass"),
    ("Usage Cap", "Exceed the daily AI-analysis cap for the current plan", "Further analyses are skipped once the cap is reached; usage count shown in Settings matches the enforced cap", "Pass"),
]
for r, row in enumerate(rows_data, start=1):
    for c, val in enumerate(row):
        fill_cell(table.rows[r].cells[c], val, center=(c == 3))

figure_caption("4.1", "Manual functional test plan and results.", space_after=14, label="Table")

body(
    "The only outstanding item at the time of writing is the Twilio-backed "
    "SMS and voice-call channels, which remain untested end-to-end because "
    "Twilio account credentials had not yet been provisioned in the "
    "deployment used for testing; the SendGrid email channel, which shares "
    "the same trigger logic and Edge Function pattern, was verified "
    "successfully, and the SMS and call channels are expected to behave "
    "identically once the corresponding secrets are configured."
)

h3("4.5 Documentation")
body(
    "Alongside the source code, the project maintains a small set of "
    "living documents that record its conventions and behaviour rather than "
    "leaving them implicit in the code alone. AGENTS.md defines the "
    "project's technology stack, development philosophy, and feature-by-"
    "feature build rules and was consulted before implementing every "
    "feature described in this chapter. IMPLEMENTATION.md records the "
    "step-by-step build plan the system was developed against, design.md "
    "records the visual design language (colour, spacing, and typography) "
    "used across the interface, and TESTING.md is the manual test checklist "
    "summarised in Table 4.1. The codebase itself is documented at the "
    "point of use: monitoring.ts, for example, opens with a comment "
    "explaining that an event is logged immediately and patched "
    "asynchronously once analysis returns, and the SQL migrations under "
    "/supabase/migrations each carry a comment block explaining the purpose "
    "of the table or function they define. This report constitutes the "
    "final layer of documentation, describing the problem, design, and "
    "implementation of the system as a whole."
)

doc.save(OUT)
print(f"Saved {OUT}")
