"""
Builds Chapter 5 (Conclusion and Recommendations) and the Appendix, and
appends them to a copy of Surveillance_AI_Chapters_1_to_4.docx.

Chapter 5 is inserted before the existing "REFERENCES" heading (so the
order becomes Chapter 4 -> Chapter 5 -> References, matching the
Bibliography-after-Chapter-5 convention in Project Format.doc). The
Appendix (A, B, C) is appended after the References list, at the very
end of the document.

Run with: cmd /c "python scripts\\build_chapter5_appendix.py"
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = "Surveillance_AI_Chapters_1_to_4.docx"
OUT = "Surveillance_AI_Full_Report.docx"

doc = Document(SRC)

# ---------------------------------------------------------------------------
# Locate the "REFERENCES" heading so Chapter 5 can be inserted before it.
# ---------------------------------------------------------------------------
references_para = None
for p in doc.paragraphs:
    if p.text.strip() == "REFERENCES" and p.style.name == "Heading 1":
        references_para = p
        break

if references_para is None:
    raise RuntimeError('Could not find a "REFERENCES" Heading 1 paragraph.')

# ---------------------------------------------------------------------------
# Helpers that insert BEFORE the References heading (used for Chapter 5).
# ---------------------------------------------------------------------------

def _new_before_refs(style):
    p = references_para.insert_paragraph_before("", style=style)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return p


def h1(text):
    p = _new_before_refs("Heading 1")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(text)
    return p


def h2(text):
    p = _new_before_refs("Heading 2")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(text)
    return p


def h3(text):
    p = _new_before_refs("Heading 3")
    p.add_run(text)
    return p


def body(text):
    p = _new_before_refs("Normal")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(10)
    p.add_run(text)
    return p


def bullet_list(items):
    for item in items:
        p = _new_before_refs("List Paragraph")
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        if isinstance(item, tuple):
            lead, rest = item
            r1 = p.add_run(f"\u2022\t{lead} ")
            r1.bold = True
            p.add_run(rest)
        else:
            p.add_run(f"\u2022\t{item}")


def numbered_list(items):
    for i, item in enumerate(items, start=1):
        p = _new_before_refs("List Paragraph")
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        if isinstance(item, tuple):
            lead, rest = item
            r1 = p.add_run(f"{i}.\t{lead} ")
            r1.bold = True
            p.add_run(rest)
        else:
            p.add_run(f"{i}.\t{item}")


# ---------------------------------------------------------------------------
# CHAPTER FIVE (inserted before References)
# ---------------------------------------------------------------------------

h1("CHAPTER FIVE")
h2("CONCLUSION AND RECOMMENDATIONS")

h3("5.1 Conclusion")
body(
    "This project set out to address a specific gap identified in Chapters "
    "One and Two: existing personal-safety applications are reactive, "
    "requiring a user to consciously recognise danger and manually trigger "
    "an alarm, they convey little contextual information when they do fire, "
    "they offer only limited passive protection for dependents, and their "
    "pricing, quotas, and underlying models are fixed at build time. "
    "Surveillance AI was designed and implemented to close this gap by "
    "making protection passive rather than reactive. Once a session is "
    "started, the system continuously and autonomously captures GPS "
    "location, camera snapshots, ambient audio, and accelerometer data, "
    "and submits the camera and audio evidence to a multimodal artificial-"
    "intelligence pipeline built on OpenAI's GPT-4o and Whisper models, "
    "producing a Low, Medium, or High risk classification on every cycle "
    "without requiring the user to do anything after the initial tap."
)
body(
    "The resulting system satisfies the functional and non-functional "
    "requirements set out in Chapter Three. A High-risk classification "
    "autonomously and simultaneously notifies a pre-configured emergency "
    "contact by SMS, email, and an automated voice call, each carrying the "
    "AI's summary and a GPS link rather than a bare distress signal, "
    "directly addressing the information-poverty problem identified in "
    "Chapter Two. Every user's data is protected by Row Level Security "
    "scoped to their own account, and every OpenAI, Twilio, and SendGrid "
    "credential is held only on the server, in a Supabase Edge Function, "
    "never on the client, satisfying the project's security objective. The "
    "three-tier plan model, together with an administrator capability to "
    "reassign a user's tier at runtime, demonstrates a concrete way in "
    "which a personal-safety product can remain configurable without a new "
    "application release, addressing the static-configuration problem "
    "identified in Chapter Two, although, as noted below, the full "
    "administrator control over pricing and quotas envisioned in Chapter "
    "One remains partly a matter for future work rather than a fully "
    "editable runtime configuration in the current build."
)
body(
    "The manual functional test plan in Table 4.1 confirmed that "
    "authentication, onboarding, session control, sensor capture, AI risk "
    "analysis, background execution, shake detection, Medium-risk "
    "notifications, High-risk email alerts, alert deduplication, stealth "
    "mode, wellness check-ins, the live map, settings persistence, degraded-"
    "permission handling, and the daily usage cap all behave as designed. "
    "The one outstanding item, end-to-end verification of the Twilio-backed "
    "SMS and voice-call channels, was not yet completed because Twilio "
    "account credentials had not been provisioned in the deployment used "
    "for testing; the parallel SendGrid email channel, which shares the "
    "same trigger logic, was verified successfully, giving confidence that "
    "the SMS and call channels will behave identically once configured. "
    "The limitations noted in Chapter One remain: the accuracy of the risk "
    "assessment is bounded by the underlying multimodal model and by "
    "prompt design, as Driessen et al. (2024) caution, and the system's "
    "usefulness depends on continuous sensor access, network connectivity, "
    "battery life, and the user's informed consent to passive monitoring of "
    "their camera and microphone. Within these limitations, the project "
    "demonstrates that a real-time, multimodal artificial-intelligence "
    "pipeline can be embedded in a mobile application to convert a "
    "conventional reactive safety alarm into a passive, autonomous, and "
    "context-aware protection system."
)

h3("5.2 Recommendations")
body(
    "The following recommendations are made for extending Surveillance AI "
    "beyond the scope of this project."
)
numbered_list([
    ("Complete Twilio provisioning.", "Provision live Twilio account credentials and repeat the High-risk SMS and voice-call test cases end to end, so that all three alert channels are verified rather than only the SendGrid email channel."),
    ("Build a true Guardian family-linking feature.", "The current Guardian tier is a higher-cap, shorter-interval subscription plan on a single account; a genuine multi-member guardian dashboard, in which one guardian account can view the status of several linked monitored users, was envisioned in Chapter One but was not implemented in this build and is recommended as the next major feature."),
    ("Implement geo-fencing.", "Add the optional geo-fence capability described in the project's early planning, so that a user leaving a defined safe area triggers a Medium-risk event independently of the AI analysis cycle."),
    ("Introduce automated testing.", "Testing in this project was entirely manual, using the checklist in Table 4.1, because the simulator cannot exercise camera, microphone, accelerometer, background execution, or push notifications; introducing unit and integration tests for the pure business logic in /lib (for example combineRisks, isCapReached, and the risk-combination pipeline) would reduce regression risk as the codebase grows."),
    ("Broaden administrator control.", "Extend the admin-users Edge Function beyond reassigning a user's existing plan tier so that an administrator can also edit each tier's price, daily cap, and minimum monitoring interval at runtime, fully realising the runtime-configurable administrative model described in Chapter One."),
    ("Extend iOS testing.", "Development and testing in this project were carried out primarily on Android; the same TypeScript codebase targets iOS through Expo, and a dedicated iOS testing pass, including Apple Sign-In, which becomes a legal requirement once any social login is offered, is recommended before any public release."),
    ("Add Google and Apple Sign-In.", "Email/password authentication was implemented for this project; social login was deferred by design and is recommended as a follow-on feature to reduce onboarding friction."),
])

# ---------------------------------------------------------------------------
# APPENDIX (appended at the very end of the document, after References)
# ---------------------------------------------------------------------------

def end_h1(text):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.add_run(text)
    return p


def end_h3(text):
    p = doc.add_paragraph(style="Heading 3")
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.add_run(text)
    return p


def end_body(text):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(10)
    p.add_run(text)
    return p


def end_bullet_list(items):
    for item in items:
        p = doc.add_paragraph(style="List Paragraph")
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.space_after = Pt(6)
        if isinstance(item, tuple):
            lead, rest = item
            r1 = p.add_run(f"\u2022\t{lead} ")
            r1.bold = True
            p.add_run(rest)
        else:
            p.add_run(f"\u2022\t{item}")


def end_numbered_list(items):
    for i, item in enumerate(items, start=1):
        p = doc.add_paragraph(style="List Paragraph")
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.space_after = Pt(6)
        if isinstance(item, tuple):
            lead, rest = item
            r1 = p.add_run(f"{i}.\t{lead} ")
            r1.bold = True
            p.add_run(rest)
        else:
            p.add_run(f"{i}.\t{item}")


def _set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "999999")
        borders.append(el)
    tblPr.append(borders)


def fill_cell(cell, text, bold=False, size=9, center=False):
    cell.paragraphs[0].text = ""
    r = cell.paragraphs[0].add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    if center:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.paragraphs[0].paragraph_format.space_after = Pt(4)


def end_table(rows, cols, widths):
    table = doc.add_table(rows=rows, cols=cols)
    _set_table_borders(table)
    for i, w in enumerate(widths):
        for row in table.rows:
            row.cells[i].width = Inches(w)
    return table


def end_caption(label, number, caption):
    cap = doc.add_paragraph(style="Normal")
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(14)
    r = cap.add_run(f"{label} {number}: {caption}")
    r.italic = True


end_h1("APPENDIX")

end_h3("Appendix A \u2013 Project Plan")
end_body(
    "The table below presents the project plan in relative weeks rather "
    "than fixed calendar dates, reflecting the order in which the system "
    "was actually built: feature by feature, starting with the backend "
    "schema and authentication and ending with manual testing and report "
    "writing, as described in Section 4.1."
)
plan_rows = [
    ("Phase 1", "Weeks 1\u20132", "Requirements gathering and literature review", "Finalised problem statement, aim, and objectives (Chapter One/Two)"),
    ("Phase 2", "Weeks 3\u20134", "System design", "Functional/non-functional requirements, architecture, and database design (Chapter Three)"),
    ("Phase 3", "Weeks 5\u20136", "Backend setup", "Supabase project, PostgreSQL schema, Row Level Security policies, email/password authentication"),
    ("Phase 4", "Weeks 7\u20138", "Core monitoring loop", "GPS/camera/audio/accelerometer capture and the analyse-image / analyse-audio Edge Functions"),
    ("Phase 5", "Weeks 9\u201310", "Alert pipeline and plan model", "send-sms / send-email / make-call Edge Functions and the Free/Pro/Guardian plan and admin model"),
    ("Phase 6", "Weeks 11\u201312", "Main interface screens", "Onboarding, Home, Live, Log, Alerts, and Settings screens"),
    ("Phase 7", "Week 13", "Testing and bug-fixing", "Manual functional test plan and results (Table 4.1)"),
    ("Phase 8", "Week 14", "Documentation and defence preparation", "Final report, diagrams, and defence slides"),
]
t = end_table(len(plan_rows) + 1, 4, [0.8, 1.1, 2.0, 2.6])
for i, htext in enumerate(["Phase", "Duration", "Activities", "Deliverable"]):
    fill_cell(t.rows[0].cells[i], htext, bold=True, center=(i < 2))
    _set_cell_shading(t.rows[0].cells[i], "D9D9D9")
for r, row in enumerate(plan_rows, start=1):
    for c, val in enumerate(row):
        fill_cell(t.rows[r].cells[c], val, center=(c < 2))
end_caption("Table", "A.1", "Project plan by development phase.")

end_h3("Appendix B \u2013 User Interfaces (more)")
end_body(
    "Section 4.3 describes the five main tab screens and stealth mode. The "
    "additional screens and states listed below are recommended for "
    "inclusion here as device screenshots, captured directly from the "
    "running application on a physical phone rather than reproduced from "
    "design mockups, so that this appendix reflects the actual delivered "
    "interface:"
)
end_bullet_list([
    "Sign Up and Sign In screens, including the invalid-credentials error state.",
    "Forgot Password and Reset Password screens.",
    "Each of the twelve onboarding screens (Landing, When, Who, Concern, Interstitial, Contact, Speed, Preferences, Social Proof, Permissions, Setup, Plan Reveal).",
    "The Home screen in both the inactive and active (session running) states, showing the shield status indicator and risk badge.",
    "The degraded-mode banner shown when camera or microphone permission is denied.",
    "The Live screen showing the GPS marker on the map.",
    "An expanded event card from the Log screen showing the full AI report, photo, and transcript.",
    "An entry on the Alerts screen showing the channels that were sent for a High-risk event.",
    "The Settings screen, including the monitoring interval and shake sensitivity pickers.",
    "The stealth-mode black overlay and the momentary reveal after a single tap.",
])

end_h3("Appendix C \u2013 User Documentation (Guide)")
end_body(
    "This appendix is a short guide to using Surveillance AI, written for a "
    "first-time user of the completed application."
)
end_numbered_list([
    ("Create an account.", "Open the app and tap Sign Up on the landing screen, enter an email address and password, and submit the form. Existing users tap Sign In instead; a Forgot Password link is available if the password is not remembered."),
    ("Complete onboarding.", "On first sign-in, the app asks a short series of questions about when and who you want protection for, then asks for your primary safety concern before collecting your emergency contact's name, phone number, and email address. It then asks you to choose a monitoring interval, shake sensitivity, and whether to enable stealth mode."),
    ("Grant permissions.", "The app requests Camera, Microphone, Location (Always), and Notifications permission one at a time, each with a short explanation of why it is needed. All four should be allowed for full protection; declining camera or microphone still allows the app to run in a degraded mode using the remaining sensors."),
    ("Start a session.", "From the Home screen, tap Start Surveillance. The shield indicator becomes active and a timer begins counting. The app will now capture your location, take periodic photos, listen to your surroundings, and watch for sudden impacts, without any further action from you."),
    ("Understand the risk badge.", "The Home and Live screens show a risk badge: green for Low, amber for Medium, and red for High. A Medium result sends you a push notification only. A High result, or a detected impact, automatically contacts your emergency contact by SMS, email, and phone call."),
    ("Check the Live map.", "Open the Live tab to see your current position on the map, along with the risk level and summary from the most recent monitoring cycle."),
    ("Review the Event Log.", "Open the Log tab to see every past monitoring cycle. Tap any entry to see its full AI summary, photo, and audio transcript where available."),
    ("Review Alerts.", "Open the Alerts tab to see a history of every High-risk event that contacted your emergency contact, including which channels were sent and when."),
    ("Stop a session.", "Tap Stop Surveillance on the Home screen at any time to immediately end monitoring."),
    ("Use stealth mode.", "If stealth mode is enabled in Settings, the screen fades to black as soon as a session starts. Tap anywhere on the black screen to reveal the interface for a few seconds; the session keeps running in the background the entire time."),
    ("Respond to a wellness check-in.", "If a check-in time is configured in Settings, the app sends a notification at that time. Tap the notification and confirm you are safe within ten minutes, or your emergency contact will be alerted automatically."),
    ("Update Settings.", "Open the Settings tab at any time to change your emergency contact, monitoring interval, shake sensitivity, stealth mode, or wellness check-in time, and to sign out."),
    ("Troubleshooting.", "If a banner appears saying camera or microphone is disabled, open your phone's system settings and re-enable the permission for the app. If the daily analysis limit for your plan is reached, the Settings screen shows your current usage; upgrading your plan raises the limit."),
])

doc.save(OUT)
print(f"Saved {OUT}")
