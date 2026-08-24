"""
Builds Chapter 3 (Design and Methodology) and appends it to a copy of the
existing Chapters 1 & 2 Word document, inserting the new chapter directly
before the existing "REFERENCES" heading so the final order is:
Chapter 1 -> Chapter 2 -> Chapter 3 -> References.

Run with: cmd /c "python scripts\build_chapter3.py"
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

SRC = "Surveillance_AI_Chapters_1_and_2_.docx"
OUT = "Surveillance_AI_Chapters_1_to_3.docx"
ASSETS = "report_assets"

doc = Document(SRC)

# ---------------------------------------------------------------------------
# Locate the "REFERENCES" heading so new content can be inserted before it.
# ---------------------------------------------------------------------------
references_para = None
for p in doc.paragraphs:
    if p.text.strip() == "REFERENCES" and p.style.name == "Heading 1":
        references_para = p
        break

if references_para is None:
    raise RuntimeError('Could not find a "REFERENCES" Heading 1 paragraph.')

# ---------------------------------------------------------------------------
# Formatting helpers replicating the exact styles/paragraph formats already
# used throughout the document (inspected from Chapters One and Two).
# ---------------------------------------------------------------------------

def _new(style):
    p = references_para.insert_paragraph_before("", style=style)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return p


def h1(text):
    p = _new("Heading 1")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(text)
    return p


def h2(text):
    p = _new("Heading 2")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(text)
    return p


def h3(text):
    p = _new("Heading 3")
    p.add_run(text)
    return p


def body(text):
    p = _new("Normal")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(10)
    p.add_run(text)
    return p


def bullet_list(items):
    """items: list of str, or (bold_lead_in, rest) tuples."""
    for item in items:
        p = _new("List Paragraph")
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        if isinstance(item, tuple):
            lead, rest = item
            r1 = p.add_run(f"\u2022\t{lead} ")
            r1.bold = True
            p.add_run(rest)
        else:
            p.add_run(f"\u2022\t{item}")


def numbered_list(items):
    for i, item in enumerate(items, start=1):
        p = _new("List Paragraph")
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        if isinstance(item, tuple):
            lead, rest = item
            r1 = p.add_run(f"{i}.\t{lead} ")
            r1.bold = True
            p.add_run(rest)
        else:
            p.add_run(f"{i}.\t{item}")


def figure(path, number, caption, width=6.3):
    p = _new("Normal")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run()
    run.add_picture(f"{ASSETS}/{path}", width=Inches(width))

    cap = _new("Normal")
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(14)
    r = cap.add_run(f"Figure {number}: {caption}")
    r.italic = True


# ---------------------------------------------------------------------------
# CHAPTER THREE
# ---------------------------------------------------------------------------

h1("CHAPTER THREE")
h2("DESIGN AND METHODOLOGY")

body(
    "This chapter presents the design and methodology of Surveillance AI. It "
    "translates the problems and conclusions established in Chapters One and "
    "Two, namely the activation gap, the information poverty of conventional "
    "alerts, the lack of passive protection for dependents, and the static, "
    "build-time configuration of existing products, into a concrete system "
    "specification and architecture. The chapter first states the proposed "
    "system and its functional and non-functional requirements, then justifies "
    "the technologies selected to build it, and finally presents the "
    "architecture and design of the system, including its database, use case "
    "and sequence diagrams."
)

h3("3.1 The Proposed System")
body(
    "Surveillance AI proposes to invert the reactive, human-initiated alarm "
    "paradigm identified in Chapter Two by making protection passive and "
    "continuous rather than dependent on a user consciously recognising "
    "danger, retrieving a locked device, and pressing a button. Once a user "
    "taps \u201cStart Surveillance\u201d, the application autonomously and "
    "repeatedly captures the user's location, a camera snapshot, ambient "
    "audio, and accelerometer readings, and submits the camera and audio "
    "evidence to a multimodal artificial-intelligence pipeline built on "
    "OpenAI's GPT-4o and Whisper models. This directly operationalises the "
    "real-time, artificial-intelligence-powered threat-detection framework "
    "argued for by Cadet et al. (2024), replacing the human operator or the "
    "user's own vigilance with an automated, continuously running risk "
    "engine. Each monitoring cycle produces a Low, Medium, or High risk "
    "classification; a Medium classification raises a push notification to "
    "the user alone, while a High classification autonomously and "
    "simultaneously notifies a pre-configured emergency contact by SMS, "
    "email, and an automated voice call, addressing the information-poverty "
    "problem by carrying the AI's summary and a GPS link rather than a bare "
    "distress signal."
)
body(
    "Because Driessen et al. (2024) caution that a multimodal model's "
    "reasoning on complex or out-of-distribution scenes remains imperfect and "
    "sensitive to prompt design, every prompt sent to GPT-4o is deliberately "
    "conservative: the model is instructed to prefer a lower risk "
    "classification unless the visual or audio evidence is unambiguous, so "
    "that the system fails toward silent logging rather than toward "
    "false alarms. The proposed system further addresses the static-"
    "configuration problem identified in Chapter Two by introducing three "
    "subscription tiers, Free, Pro, and Guardian, each with its own daily "
    "artificial-intelligence usage cap, minimum monitoring interval, and "
    "audio-analysis entitlement, together with an administrator capability "
    "to reassign a user's tier at runtime without a new application release. "
    "The Guardian tier, in particular, is designed to extend the benefit of "
    "the system to dependents such as children, elderly relatives, and "
    "other family members who may be less able to manage a safety "
    "application unaided."
)

h3("3.2 System Specification")

h3("3.2.1 Functional Requirements")
body(
    "The functional requirements below describe the observable behaviour "
    "that the implemented system provides to its users and administrator."
)
numbered_list([
    ("Registration and Authentication.", "The system shall allow a user to register and sign in with an email address and password through Supabase Authentication, and to request a password reset via a deep link that reopens the application."),
    ("Guided Onboarding.", "The system shall guide a first-time user through a fixed sequence of screens that collect the emergency contact's name, phone number, and email address, and store this information securely for use during an alert."),
    ("Configurable Preferences.", "The system shall allow the user to configure the monitoring interval (20, 30, or 60 seconds), shake sensitivity (Low, Medium, or High), stealth mode, and a wellness check-in time."),
    ("One-Tap Session Control.", "The system shall allow the user to start and stop a surveillance session with a single tap, creating and later closing a corresponding session record."),
    ("Parallel Sensor Capture.", "While a session is active, the system shall capture the device's GPS location, a silent camera snapshot, and an ambient audio clip in parallel on every monitoring cycle."),
    ("Continuous Shake Detection.", "The system shall continuously monitor accelerometer data, independently of the timed monitoring cycle, and classify a sudden qualifying impact as a High-risk event immediately."),
    ("Image Risk Analysis.", "The system shall submit each captured photograph, together with the current GPS coordinates, to the analyse-image function for GPT-4o Vision risk analysis."),
    ("Audio Risk Analysis.", "The system shall transcribe each captured audio clip using Whisper and submit the resulting transcript to the analyse-audio function for GPT-4o threat analysis, without blocking the delivery of the image-based result."),
    ("Consolidated Risk Scoring.", "The system shall combine the image-based and audio-based risk outputs of a monitoring cycle into a single consolidated risk level of Low, Medium, or High."),
    ("Medium-Risk Notification.", "On a Medium-risk classification, the system shall deliver a local push notification to the user only."),
    ("High-Risk Alert Pipeline.", "On a High-risk classification, the system shall automatically and simultaneously send an SMS and an email, and place an automated voice call, to the registered emergency contact, each carrying the GPS location, the AI-generated summary, and the photograph where available, and shall not send a duplicate alert for the same event."),
    ("Event Log and Alert History.", "The system shall record every monitoring cycle to an Event Log, showing its timestamp, risk level, AI summary, and coordinates, and shall record every High-risk alert to a separate Alerts screen showing which channels were sent and when."),
    ("Live Location Map.", "The system shall display the user's current location on an interactive map during an active session."),
    ("Tiered Usage Plans.", "The system shall enforce a per-user, per-day cap on artificial-intelligence analysis calls according to the user's subscription plan (Free, Pro, or Guardian), and shall allow an authorised administrator to list registered users and reassign a user's plan at runtime."),
    ("Wellness Check-In.", "The system shall raise a scheduled notification at the user's configured check-in time and, if the user has not confirmed safety within ten minutes, shall alert the emergency contact."),
])

h3("3.2.2 Non-Functional Requirements")
body(
    "The non-functional requirements below describe the quality attributes "
    "that constrain how the functional requirements are satisfied."
)
bullet_list([
    ("Security.", "Every database table has Row Level Security enabled and scoped to auth.uid(), so a user can read and write only their own data; all OpenAI, Twilio, and SendGrid API keys are stored as Supabase secrets and are never bundled into the client, with every third-party call proxied through a Supabase Edge Function."),
    ("Performance.", "GPS, camera, and audio capture run in parallel within a monitoring cycle rather than sequentially, and audio transcription and analysis run asynchronously so that a slower audio pipeline never delays the delivery of the image-based risk result."),
    ("Reliability.", "The monitoring loop is registered with Expo Task Manager and Background Fetch so that cycles continue while the application is minimised or the screen is locked, and every Edge Function returns a safe low-risk fallback result rather than throwing an error if the underlying AI call fails."),
    ("Usability.", "Camera, microphone, location, and notification permissions are requested one at a time during onboarding, each preceded by a plain-language explanation of why it is needed, and a denied camera or microphone permission degrades the experience with a visible banner rather than crashing the application."),
    ("Scalability.", "The system is built on Supabase's managed PostgreSQL and Edge Function infrastructure, and per-tier daily usage caps bound the growth of third-party AI cost as the number of users increases."),
    ("Maintainability.", "All AI prompts are centralised in a single prompts.ts file and all plan parameters (price, daily cap, minimum interval, audio entitlement) are centralised in a single plans.ts file, so behaviour can be adjusted in one place rather than across the codebase."),
    ("Portability.", "The application is built with Expo and React Native so that a single TypeScript codebase targets both Android and iOS; development and testing in this project were carried out primarily on Android."),
    ("Privacy.", "Camera and microphone data are used only to produce a risk assessment for the active session, and the Event Log is automatically cleared after five days by default, limiting how long sensitive personal data is retained on the device."),
])

h3("3.3 Selection of Technologies and Tools")
body(
    "The technologies selected for Surveillance AI were chosen to keep a "
    "single TypeScript codebase, to avoid exposing third-party credentials on "
    "the client, and to minimise the amount of custom backend infrastructure "
    "the project would otherwise need to build and operate."
)
bullet_list([
    ("Expo, React Native, and TypeScript.", "provide a single, statically typed codebase that targets both Android and iOS and give access to native device capabilities (camera, location, sensors, background tasks) through a consistent, well-documented API."),
    ("Expo Router.", "supplies file-based navigation for the application's authentication, onboarding, and tab-based screens."),
    ("NativeWind (Tailwind CSS).", "allows the interface to be styled with utility classes directly in component code, keeping presentation close to markup."),
    ("Zustand and AsyncStorage.", "manage in-memory application state (session, alerts, settings, onboarding) and persist it locally so the app remains responsive and usable offline between synchronisations."),
    ("Supabase (PostgreSQL, Authentication, Row Level Security, Edge Functions).", "supplies a managed relational database, email/password authentication, and a serverless runtime, removing the need to design and operate a bespoke server while still allowing every third-party API key to be kept off the client."),
    ("OpenAI GPT-4o and Whisper.", "provide, respectively, the multimodal vision-and-language reasoning used to interpret camera snapshots and the speech-to-text transcription used to make ambient audio analysable by GPT-4o."),
    ("Twilio.", "delivers the SMS and automated voice call channels of the High-risk alert pipeline."),
    ("SendGrid.", "delivers the email channel of the High-risk alert pipeline."),
    ("Expo Camera, Expo AV, Expo Location, and Expo Sensors.", "provide, respectively, silent photo capture, audio recording, GPS access, and accelerometer readings for shake detection."),
    ("Expo Task Manager and Background Fetch.", "register and continue the monitoring loop while the application is minimised or the screen is locked."),
    ("Expo Secure Store, Keep Awake, Notifications, and SMS.", "provide, respectively, on-device storage of the emergency contact, prevention of screen sleep during stealth mode, local and push notifications, and a fallback SMS composer."),
    ("react-native-webview and Leaflet.", "render the Live Map using CartoDB Dark Matter tiles inside a self-contained HTML page, avoiding a native maps SDK and its associated API key."),
])

h3("3.4 Architecture of the System")
body(
    "Surveillance AI follows a three-tier architecture comprising the mobile "
    "client, the Supabase backend, and a set of external third-party "
    "services, as shown in Figure 3.1. The mobile client owns the user "
    "interface, local state, and the monitoring orchestrator that drives "
    "each capture cycle. It communicates with Supabase for authentication "
    "and for synchronising sessions, events, alerts, contacts, and settings, "
    "all of which are protected by Row Level Security. Rather than call "
    "OpenAI, Twilio, or SendGrid directly, the client invokes a small set of "
    "Supabase Edge Functions (analyse-image, analyse-audio, send-sms, "
    "send-email, make-call, and admin-users/upgrade-plan), each of which "
    "holds the corresponding third-party credential as a server-side secret "
    "and performs the actual third-party call on the client's behalf. This "
    "arrangement means that no OpenAI, Twilio, or SendGrid key is ever "
    "present in the mobile application binary."
)
figure("architecture.png", "3.1", "Three-tier system architecture of Surveillance AI.")

h3("3.5 Design of the System")
body(
    "Two complementary diagrams are presented to describe the design of the "
    "system: a use case diagram, which identifies the actors of the system "
    "and the functions each actor performs, and a sequence diagram, which "
    "traces the order of interactions that occur during a single monitoring "
    "cycle and the alert pipeline it may trigger."
)

h3("3.5.1 Use Case and Sequence Diagrams")
body(
    "Figure 3.2 presents the use case diagram of the system. The Registered "
    "User is the principal actor and interacts with the system to register "
    "and sign in, complete onboarding and set an emergency contact, start "
    "and stop a surveillance session, view the live map and event log, "
    "configure monitoring preferences, and perform a wellness check-in. "
    "Starting a session triggers the monitoring cycle use case, which in "
    "turn may result in the user receiving a Medium-risk push notification "
    "or a High-risk alert being sent to the emergency contact. A second "
    "actor, the Administrator, interacts with the system separately to "
    "manage user plans and usage quotas."
)
figure("usecase.png", "3.2", "Use case diagram of Surveillance AI.")

body(
    "Figure 3.3 presents the sequence diagram of a single monitoring cycle. "
    "After the user starts a session, the orchestrator repeats a fixed loop "
    "roughly every twenty to thirty seconds: it captures the GPS location "
    "and a camera snapshot in parallel, submits the photograph and location "
    "to the analyse-image Edge Function for GPT-4o Vision analysis, and, "
    "without waiting for that result, records and submits an ambient audio "
    "clip to the analyse-audio Edge Function for Whisper transcription and "
    "GPT-4o threat analysis. The image-based and audio-based risk results "
    "are combined into one risk level for the cycle. A Medium result "
    "produces a local push notification to the user only, while a High "
    "result causes the orchestrator to invoke the send-sms, send-email, and "
    "make-call Edge Functions, which dispatch the alert to the emergency "
    "contact through Twilio and SendGrid."
)
figure("sequence.png", "3.3", "Sequence diagram of a monitoring cycle and the High-risk alert pipeline.")

h3("3.5.2 Database Design")
body(
    "The system's persistent data is held in a PostgreSQL database managed "
    "by Supabase and organised into seven tables, all protected by Row "
    "Level Security policies that restrict every row to the authenticated "
    "user identified by auth.uid(). The entity-relationship diagram in "
    "Figure 3.4 shows these tables and the foreign-key relationships "
    "between them."
)
bullet_list([
    ("users.", "one row per authenticated user, created automatically by a database trigger on first sign-up, storing the user's subscription plan (free, pro, or guardian) and account creation timestamp."),
    ("sessions.", "one row per surveillance session, recording when the session started and ended and how many monitoring cycles it completed."),
    ("events.", "one row per monitoring cycle, recording its timestamp, consolidated risk level, AI-generated summary, GPS coordinates, and, where available, a photo URL and audio transcript."),
    ("alerts.", "created only for High-risk events, recording which channels (SMS, email, call) were sent, to which contact, and when, and used to prevent the same event from triggering a duplicate alert."),
    ("contacts.", "the emergency contact's name, phone number, and email address collected during onboarding or updated in Settings."),
    ("settings.", "one row per user, storing the monitoring interval, shake sensitivity, stealth mode toggle, and wellness check-in time."),
    ("ai_usage.", "one row per user per calendar date, incremented atomically by a database function to enforce each plan tier's daily artificial-intelligence usage cap."),
])
figure("erd.png", "3.4", "Entity-relationship diagram of the Surveillance AI database.")

doc.save(OUT)
print(f"Saved {OUT}")
