"""
Adds a large second batch of real code excerpts to Chapter 4, Section 4.2
(Programming and Coding) of Surveillance_AI_Full_Report.docx, targeting
roughly 10 additional pages of code content as requested. Inserts directly
before the "4.3 Main User Interfaces" heading so nothing after it (the
user's screenshots) is touched. A .backup2.docx copy was made first.

Run with: cmd /c "python scripts\\add_more_code_ch4_v2.py"
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

TARGET = "Surveillance_AI_Full_Report.docx"

doc = Document(TARGET)

anchor = None
for p in doc.paragraphs:
    if p.text.strip() == "4.3 Main User Interfaces" and p.style.name == "Heading 3":
        anchor = p
        break

if anchor is None:
    raise RuntimeError('Could not find the "4.3 Main User Interfaces" heading.')


def _new(style):
    p = anchor.insert_paragraph_before("", style=style)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return p


def body(text):
    p = _new("Normal")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(10)
    p.add_run(text)
    return p


def code_block(lines):
    for line in lines:
        p = _new("Normal")
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.3)
        r = p.add_run(line if line else "\u00A0")
        r.font.name = "Consolas"
        r.font.size = Pt(9)


def figure_caption(number, caption, space_after=14):
    cap = _new("Normal")
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(space_after)
    r = cap.add_run(f"Figure {number}: {caption}")
    r.italic = True


# ---------------------------------------------------------------------------
# Batch 2 — eleven further code excerpts for Section 4.2
# ---------------------------------------------------------------------------

body(
    "The remaining figures in this section extend the walkthrough to the "
    "audio-analysis path, the client and server sides of the AI integration, "
    "the GPS and camera capture utilities, push notifications, the two "
    "Zustand stores that hold session and event state, the Twilio SMS Edge "
    "Function, the hand-written Leaflet map, and the Row Level Security "
    "policies that protect every table."
)

body(
    "Audio transcription is never awaited inside the main capture path. "
    "Figure 4.7 shows how, once the audio clip finishes recording, the "
    "transcript and its risk score are combined with the image risk already "
    "computed, and the combined result is only allowed to escalate a level "
    "that has not already triggered a notification, preventing the same "
    "cycle from firing two notifications at different points in time."
)
code_block([
    "const [audioUri, imageResult] = await Promise.all([",
    "  audioPromise.catch((err) => { console.error(err); return null; }),",
    "  imagePromise.catch(() => null),",
    "]);",
    "",
    "const imageRisk: RiskLevel = imageResult?.riskLevel ?? \"low\";",
    "",
    "if (audioUri) {",
    "  useAlertStore.getState().updateEvent(event.id, { audioUri });",
    "  transcribeAudio(audioUri).then(async (transcript) => {",
    "    if (!transcript) return;",
    "    useAlertStore.getState().updateEvent(event.id, { transcript });",
    "",
    "    const { audioSummary, audioRisk } = await analyseAudioTranscript(transcript);",
    "    const consolidatedRisk = combineRisks(imageRisk, audioRisk);",
    "",
    "    // Only escalate — never double-alert a level already fired above.",
    "    if (consolidatedRisk === \"high\" && imageRisk !== \"high\") {",
    "      sendLocalNotification(\"HIGH RISK DETECTED\", audioSummary ?? \"\");",
    "    }",
    "  });",
    "}",
])
figure_caption("4.7", "Non-blocking audio transcription and risk escalation logic (monitoring.ts).")

body(
    "On the client, every Edge Function call goes through the Supabase "
    "functions client rather than a raw fetch, so that the user's auth "
    "session is attached automatically. Figure 4.8 shows analyseImage, "
    "which additionally unwraps the FunctionsHttpError's underlying "
    "Response so that the real server-side error, such as a missing "
    "OpenAI key or a 4xx from OpenAI, is logged instead of the generic "
    "\u201cnon-2xx status code\u201d message Supabase returns by default."
)
code_block([
    "export async function analyseImage(",
    "  imageBase64: string,",
    "  location: { lat: number; lng: number } | null,",
    "): Promise<VisionResult | null> {",
    "  const { data, error } = await supabase.functions.invoke(\"analyse-image\", {",
    "    body: { imageBase64, lat: location?.lat ?? null, lng: location?.lng ?? null },",
    "  });",
    "",
    "  if (error) {",
    "    const ctx = (error as unknown as { context?: Response }).context;",
    "    if (ctx) {",
    "      const body = await ctx.text().catch(() => \"(unreadable)\");",
    "      console.error(`[vision] analyseImage error: HTTP ${ctx.status} —`, body);",
    "    } else {",
    "      console.error(\"[vision] analyseImage error:\", error.message);",
    "    }",
    "    return null;",
    "  }",
    "",
    "  return data as VisionResult;",
    "}",
])
figure_caption("4.8", "Client-side Edge Function invocation with detailed error extraction (vision.ts).")

body(
    "Audio must be recorded in a Whisper-compatible container, which "
    "differs by platform. Figure 4.9 shows the recording configuration "
    "used by expo-audio: an uncompressed WAV container on iOS, because "
    "AVAudioRecorder infers the container from the file extension and "
    "compressed formats were found to produce an unreadable file, and an "
    "AAC-in-MP4 container on Android, which MediaRecorder writes reliably."
)
code_block([
    "const WHISPER_RECORDING_OPTIONS: RecordingOptions = {",
    "  extension: Platform.OS === \"ios\" ? \".wav\" : \".mp4\",",
    "  sampleRate: 16000,",
    "  numberOfChannels: 1,",
    "  bitRate: 128000,",
    "  isMeteringEnabled: false,",
    "  ios: {",
    "    outputFormat: IOSOutputFormat.LINEARPCM,",
    "    audioQuality: AudioQuality.HIGH,",
    "    linearPCMBitDepth: 16,",
    "    linearPCMIsBigEndian: false,",
    "    linearPCMIsFloat: false,",
    "  },",
    "  android: {",
    "    extension: \".mp4\",",
    "    outputFormat: \"mpeg4\",",
    "    audioEncoder: \"aac\",",
    "  },",
    "  web: { mimeType: \"audio/webm\", bitsPerSecond: 64000 },",
    "};",
])
figure_caption("4.9", "Platform-specific, Whisper-compatible audio recording configuration (audio.ts).")

body(
    "Reverse geocoding turns a raw GPS coordinate into a human-readable "
    "address for the event log and SMS alert. Figure 4.10 shows how this "
    "call is raced against a three-second timeout, because Android's "
    "native geocoder can otherwise block for the full five-second "
    "operating-system timeout and hold up the monitoring cycle."
)
code_block([
    "export async function reverseGeocode(",
    "  lat: number,",
    "  lng: number,",
    "): Promise<Address | null> {",
    "  let timer: ReturnType<typeof setTimeout> | undefined;",
    "  try {",
    "    const timeout = new Promise<never>((_, reject) => {",
    "      timer = setTimeout(() => reject(new Error(\"reverseGeocode timeout\")), 3000);",
    "    });",
    "    const results = await Promise.race([",
    "      Location.reverseGeocodeAsync({ latitude: lat, longitude: lng }),",
    "      timeout,",
    "    ]);",
    "    if (!results.length) return null;",
    "    const r = results[0];",
    "    return { name: r.name ?? null, street: r.street ?? null, city: r.city ?? null,",
    "              region: r.region ?? null, country: r.country ?? null };",
    "  } catch {",
    "    return null;",
    "  } finally {",
    "    clearTimeout(timer);",
    "  }",
    "}",
])
figure_caption("4.10", "Timeout-guarded reverse geocoding (location.ts).")

body(
    "Because expo-camera has no headless capture API, a picture can only "
    "be taken from a mounted CameraView. Figure 4.11 shows the ref-"
    "registration pattern that lets a one-pixel, invisible CameraView "
    "rendered by SilentCamera.tsx be driven from plain functions, so the "
    "rest of the codebase can call takeSnapshot() without knowing that a "
    "camera component exists at all."
)
code_block([
    "let _ref: RefObject<CameraView | null> | null = null;",
    "let _ready = false;",
    "",
    "export function registerCameraRef(ref: RefObject<CameraView | null>) {",
    "  _ref = ref;",
    "  _ready = false;",
    "}",
    "",
    "export function setCameraReady() {",
    "  _ready = true;",
    "}",
    "",
    "export function clearCameraRef() {",
    "  _ref = null;",
    "  _ready = false;",
    "}",
    "",
    "export async function takeSnapshot(): Promise<string | null> {",
    "  if (!_ref?.current || !_ready) return null;",
    "  const photo = await _ref.current.takePictureAsync({ quality: 0.5 });",
    "  return photo?.uri ?? null;",
    "}",
])
figure_caption("4.11", "Headless snapshot capture via a registered CameraView ref (camera.ts).")

body(
    "Medium-risk events raise a local push notification rather than an "
    "alert to the emergency contact. Figure 4.12 shows the permission "
    "request together with the Android notification channel that must be "
    "created before any local alert can appear on that platform, since "
    "Android silently drops notifications sent to an unregistered channel."
)
code_block([
    "export async function requestNotificationPermission(): Promise<boolean> {",
    "  const { status: existing } = await Notifications.getPermissionsAsync();",
    "  if (existing === \"granted\") return true;",
    "  const { status } = await Notifications.requestPermissionsAsync();",
    "  return status === \"granted\";",
    "}",
    "",
    "export async function registerForPushNotifications(): Promise<string | null> {",
    "  if (!Device.isDevice) return null;",
    "  const granted = await requestNotificationPermission();",
    "  if (!granted) return null;",
    "",
    "  if (Platform.OS === \"android\") {",
    "    await Notifications.setNotificationChannelAsync(\"safety-alert\", {",
    "      name: \"Safety Alerts\",",
    "      importance: Notifications.AndroidImportance.MAX,",
    "      vibrationPattern: [0, 250, 250, 250],",
    "      lightColor: \"#00E5FF\",",
    "    });",
    "  }",
    "  return null;",
    "}",
])
figure_caption("4.12", "Notification permission request and Android channel setup (notifications.ts).")

body(
    "Session state is held in a persisted Zustand store so that an active "
    "session survives an app restart. Figure 4.13 shows the slice that "
    "starts and stops a session: a new session ID is generated locally "
    "from the current timestamp, avoiding a network round-trip before "
    "monitoring can begin, and stopping a session resets every transient "
    "field back to its idle state."
)
code_block([
    "startSession: () =>",
    "  set({",
    "    isActive: true,",
    "    sessionId: Date.now().toString(),",
    "    sessionStartTime: Date.now(),",
    "    lastRiskLevel: \"low\",",
    "    cycleCount: 0,",
    "  }),",
    "",
    "stopSession: () =>",
    "  set({",
    "    isActive: false,",
    "    sessionId: null,",
    "    sessionStartTime: null,",
    "    lastRiskLevel: null,",
    "    lastAISummary: null,",
    "    cycleCount: 0,",
    "  }),",
])
figure_caption("4.13", "Persisted session-lifecycle slice (useSessionStore.ts).")

body(
    "The alert store keeps a local, optimistic copy of every fired alert "
    "and syncs it to Supabase in the background so the Alerts screen never "
    "waits on the network. Figure 4.14 shows addAlert, which deduplicates "
    "by event ID before doing anything else, updates local state "
    "immediately, and then inserts the row into Supabase without rolling "
    "back the local copy if the network write fails."
)
code_block([
    "addAlert: async (alert) => {",
    "  // Deduplicate: skip if an alert for this event already exists.",
    "  if (get().alerts.some((a) => a.eventId === alert.eventId)) return;",
    "",
    "  // Optimistic local update — Alerts screen updates instantly.",
    "  set((state) => ({ alerts: [alert, ...state.alerts].slice(0, MAX_EVENTS) }));",
    "",
    "  const userId = useSessionStore.getState().userId;",
    "  if (!userId) return;",
    "",
    "  const { error } = await supabase.from(\"alerts\").insert({",
    "    id: alert.id,",
    "    event_id: alert.eventId,",
    "    user_id: userId,",
    "    contact_name: alert.contactName,",
    "    sms_sent: alert.smsSent,",
    "    email_sent: alert.emailSent,",
    "    call_made: alert.callMade,",
    "  });",
    "",
    "  if (error) {",
    "    console.error(\"[useAlertStore] Failed to sync alert:\", error.message);",
    "    // Alert stays in local AsyncStorage — do not roll back.",
    "  }",
    "},",
])
figure_caption("4.14", "Optimistic local update with background Supabase sync (useAlertStore.ts).")

body(
    "The SMS channel of the alert pipeline is implemented as its own "
    "Supabase Edge Function so that the Twilio credentials never reach "
    "the client. Figure 4.15 shows the full function: it validates the "
    "request body, reads the three Twilio secrets from the Edge Function "
    "environment, and posts the message to the Twilio REST API using HTTP "
    "basic authentication built from the account SID and auth token."
)
code_block([
    "Deno.serve(async (req) => {",
    "  if (req.method === \"OPTIONS\") return new Response(\"ok\", { headers: CORS });",
    "",
    "  const { to, message } = await req.json();",
    "  if (!to || !message) {",
    "    return json({ success: false, error: \"to and message are required\" }, 400);",
    "  }",
    "",
    "  const accountSid = Deno.env.get(\"TWILIO_ACCOUNT_SID\");",
    "  const authToken = Deno.env.get(\"TWILIO_AUTH_TOKEN\");",
    "  const from = Deno.env.get(\"TWILIO_PHONE_NUMBER\");",
    "  if (!accountSid || !authToken || !from) {",
    "    return json({ success: false, error: \"Twilio not configured\" }, 500);",
    "  }",
    "",
    "  const credentials = btoa(`${accountSid}:${authToken}`);",
    "  const url = `https://api.twilio.com/2010-04-01/Accounts/${accountSid}/Messages.json`;",
    "  const body = new URLSearchParams({ To: to, From: from, Body: message });",
    "",
    "  const response = await fetch(url, {",
    "    method: \"POST\",",
    "    headers: {",
    "      Authorization: `Basic ${credentials}`,",
    "      \"Content-Type\": \"application/x-www-form-urlencoded\",",
    "    },",
    "    body: body.toString(),",
    "  });",
    "",
    "  if (!response.ok) return json({ success: false, error: \"Twilio request failed\" }, 502);",
    "  return json({ success: true });",
    "});",
])
figure_caption("4.15", "Twilio SMS dispatch Edge Function (send-sms/index.ts).")

body(
    "The Live map is not built with a native maps library, since the "
    "project deliberately renders CartoDB Dark Matter tiles inside a "
    "WebView using a small, self-contained Leaflet-style HTML page. "
    "Figure 4.16 shows the coordinate-to-tile conversion functions that "
    "the page uses to work out which tile images to request for the "
    "user's current latitude and longitude at a given zoom level."
)
code_block([
    "var TILE = 256;",
    "function lng2x(ln, z) {",
    "  return (ln + 180) / 360 * Math.pow(2, z) * TILE;",
    "}",
    "function lat2y(la, z) {",
    "  var r = la * Math.PI / 180;",
    "  return (1 - Math.log(Math.tan(r) + 1 / Math.cos(r)) / Math.PI) / 2",
    "         * Math.pow(2, z) * TILE;",
    "}",
    "function x2lng(x, z) {",
    "  return x / (Math.pow(2, z) * TILE) * 360 - 180;",
    "}",
    "function y2lat(y, z) {",
    "  return Math.atan(Math.sinh(Math.PI * (1 - 2 * y / (Math.pow(2, z) * TILE))))",
    "         * 180 / Math.PI;",
    "}",
])
figure_caption("4.16", "Latitude/longitude to slippy-map tile coordinate conversion (LiveMap.tsx).")

body(
    "Every table in the database enforces Row Level Security so that a "
    "user can only see and modify their own rows. Figure 4.17 shows the "
    "recurring policy pattern, applied here to the sessions table: select, "
    "insert, update, and delete are each gated by a separate policy that "
    "compares the row's user_id column against auth.uid(), the ID of the "
    "currently authenticated Supabase user making the request."
)
code_block([
    "create table if not exists sessions (",
    "  id           uuid        primary key default gen_random_uuid(),",
    "  user_id      uuid        not null references users (id) on delete cascade,",
    "  started_at   timestamptz not null default now(),",
    "  ended_at     timestamptz,",
    "  total_cycles integer     not null default 0",
    ");",
    "",
    "alter table sessions enable row level security;",
    "",
    "create policy \"sessions: select own rows\"",
    "  on sessions for select",
    "  using ((select auth.uid()) = user_id);",
    "",
    "create policy \"sessions: insert own rows\"",
    "  on sessions for insert",
    "  with check ((select auth.uid()) = user_id);",
    "",
    "create policy \"sessions: update own rows\"",
    "  on sessions for update",
    "  using ((select auth.uid()) = user_id);",
    "",
    "create policy \"sessions: delete own rows\"",
    "  on sessions for delete",
    "  using ((select auth.uid()) = user_id);",
])
figure_caption("4.17", "Row Level Security policy pattern applied to the sessions table (001_initial_schema.sql).")

doc.save(TARGET)
print(f"Updated {TARGET} in place (backup saved as Surveillance_AI_Full_Report.backup2.docx)")
