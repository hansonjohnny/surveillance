"""
Adds more real code excerpts to Chapter 4, Section 4.2 (Programming and
Coding) of the CURRENT Surveillance_AI_Full_Report.docx, which already
contains the user's manually inserted screenshots. Operates in place on
that file (a .backup.docx copy was made first) and inserts new content
directly before the "4.3 Main User Interfaces" heading, so nothing after
that point (including the user's screenshots) is touched.

Run with: cmd /c "python scripts\\add_more_code_ch4.py"
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

TARGET = "Surveillance_AI_Full_Report.docx"

doc = Document(TARGET)

# ---------------------------------------------------------------------------
# Locate the "4.3 Main User Interfaces" heading — new content goes before it.
# ---------------------------------------------------------------------------
anchor = None
for p in doc.paragraphs:
    if p.text.strip() == "4.3 Main User Interfaces" and p.style.name == "Heading 3":
        anchor = p
        break

if anchor is None:
    raise RuntimeError('Could not find the "4.3 Main User Interfaces" heading.')


def _new(style):
    p = anchor.insert_paragraph_before("", style=style)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return p


def body(text):
    p = _new("Normal")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(10)
    p.add_run(text)
    return p


def code_block(lines):
    for line in lines:
        p = _new("Normal")
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.3)
        r = p.add_run(line if line else "\u00A0")
        r.font.name = "Consolas"
        r.font.size = Pt(9)


def figure_caption(number, caption, space_after=14):
    cap = _new("Normal")
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(space_after)
    r = cap.add_run(f"Figure {number}: {caption}")
    r.italic = True


# ---------------------------------------------------------------------------
# Additional code excerpts for Section 4.2
# ---------------------------------------------------------------------------

body(
    "The excerpts above show the client-side capture cycle and the "
    "database's atomic usage-cap function. The remaining figures in this "
    "section illustrate four further representative parts of the "
    "implementation: continuous shake detection, the server-side call to "
    "GPT-4o Vision with its safe-fallback handling, the alert pipeline's "
    "deduplication and message composition, and the registration of the "
    "background task that drives the wellness check-in."
)

body(
    "Shake detection runs as its own continuous accelerometer listener, "
    "entirely independent of the timed monitoring cycle, so that a violent "
    "impact is classified as High risk within about half a second rather "
    "than waiting for the next scheduled cycle. Figure 4.3 shows the "
    "magnitude threshold check and the 500-millisecond sustained-impact and "
    "5-second cooldown logic that together prevent a single fall from "
    "firing repeated alerts as the phone continues to move after impact."
)
code_block([
    "const subscription = Accelerometer.addListener(({ x, y, z }) => {",
    "  const magnitude = Math.sqrt(x * x + y * y + z * z);",
    "",
    "  if (magnitude > threshold) {",
    "    if (thresholdExceededAt === null) thresholdExceededAt = Date.now();",
    "",
    "    const sustainedFor = Date.now() - thresholdExceededAt;",
    "    const cooldownPassed = Date.now() - lastShakeAt > 5000;",
    "",
    "    if (sustainedFor >= 500 && cooldownPassed) {",
    "      lastShakeAt = Date.now();",
    "      thresholdExceededAt = null;",
    "      onShake();",
    "    }",
    "  } else {",
    "    // Reset unless the impact is sustained continuously, not cumulatively.",
    "    thresholdExceededAt = null;",
    "  }",
    "});",
])
figure_caption("4.3", "Continuous, threshold-based shake detection with cooldown (sensors.ts).")

body(
    "The analyse-image Edge Function is the only part of the system that "
    "holds the OpenAI API key, and it calls GPT-4o Vision with the captured "
    "photograph, encoded as a base64 data URL, alongside the current GPS "
    "coordinates. Figure 4.4 shows the request together with the safe-"
    "fallback branch: if OpenAI responds with anything other than a "
    "successful status, or if the call throws, the function returns a "
    "well-formed Low-risk result instead of an error, which is why a "
    "temporary OpenAI outage produces an \u201cAnalysis failed\u201d log "
    "entry rather than a crash."
)
code_block([
    "const response = await fetch(\"https://api.openai.com/v1/chat/completions\", {",
    "  method: \"POST\",",
    "  headers: { Authorization: `Bearer ${OPENAI_API_KEY}`, ... },",
    "  body: JSON.stringify({",
    "    model: \"gpt-4o\",",
    "    max_tokens: 256,",
    "    messages: [{",
    "      role: \"user\",",
    "      content: [",
    "        { type: \"image_url\", image_url: { url: dataUrl, detail: \"low\" } },",
    "        { type: \"text\", text: buildPrompt(lat, lng, situation) },",
    "      ],",
    "    }],",
    "  }),",
    "});",
    "",
    "if (!response.ok) {",
    "  console.error(\"[analyse-image] OpenAI error:\", response.status);",
    "  return json({ riskLevel: \"low\", summary: \"Analysis failed.\", confidence: 0 });",
    "}",
])
figure_caption("4.4", "GPT-4o Vision request with safe-fallback error handling (analyse-image Edge Function).")

body(
    "The alert orchestrator in alerts.ts is responsible for both preventing "
    "duplicate alerts and composing the message sent to the emergency "
    "contact. Figure 4.5 shows the deduplication check, which looks up the "
    "triggering event's ID against every alert already recorded before any "
    "channel is dispatched, together with the SMS message template that "
    "carries the AI summary, a Google Maps link, and the event timestamp."
)
code_block([
    "// Deduplicate \u2014 if an alert for this event was already fired, bail out.",
    "const existingAlerts = useAlertStore.getState().alerts;",
    "if (existingAlerts.some((a) => a.eventId === event.id)) {",
    "  console.log(\"[alerts] duplicate event, skipping:\", event.id);",
    "  return;",
    "}",
    "",
    "const mapsLink = event.location",
    "  ? `https://maps.google.com/?q=${event.location.lat},${event.location.lng}`",
    "  : \"Location unavailable\";",
    "",
    "const smsMessage =",
    "  `SAFETY ALERT: ${contact.name} may need help.\\n` +",
    "  `Map: ${mapsLink}\\n` +",
    "  `AI detected: ${event.aiSummary}\\n` +",
    "  `Time: ${timestamp}`;",
])
figure_caption("4.5", "Alert deduplication and SMS message composition (alerts.ts).")

body(
    "Finally, the wellness check-in feature relies on Expo Task Manager and "
    "Background Fetch to keep running even when the application itself is "
    "not open. Figure 4.6 shows the task definition, which must be "
    "registered at module load time as required by Task Manager, and the "
    "registration call, which is safe to invoke on every app launch because "
    "it checks whether the task is already registered before scheduling it "
    "again."
)
code_block([
    "// Defined at module load time \u2014 required by TaskManager.",
    "TaskManager.defineTask(WELLNESS_CHECK_TASK, async () => {",
    "  const shouldAlert = await checkWellnessWindow();",
    "  if (!shouldAlert) return BackgroundFetch.BackgroundFetchResult.NoData;",
    "",
    "  await triggerWellnessAlert(contact);",
    "  return BackgroundFetch.BackgroundFetchResult.NewData;",
    "});",
    "",
    "export async function registerWellnessTask(): Promise<void> {",
    "  const isRegistered = await TaskManager.isTaskRegisteredAsync(WELLNESS_CHECK_TASK);",
    "  if (isRegistered) return;",
    "",
    "  await BackgroundFetch.registerTaskAsync(WELLNESS_CHECK_TASK, {",
    "    minimumInterval: 10 * 60,",
    "    stopOnTerminate: false,",
    "    startOnBoot: true,",
    "  });",
    "}",
])
figure_caption("4.6", "Background task definition and registration for the wellness check-in (wellnessTask.ts).")

doc.save(TARGET)
print(f"Updated {TARGET} in place (backup saved as Surveillance_AI_Full_Report.backup.docx)")
